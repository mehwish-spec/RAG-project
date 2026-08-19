import io

import docx
import pytest
from pypdf import PdfWriter

from app.core.exceptions import ExtractionError, NoExtractableTextError
from app.ingestion.docx_loader import DOCXLoader
from app.ingestion.pdf_loader import PDFLoader
from app.ingestion.registry import get_loader
from app.ingestion.txt_loader import TXTLoader


def _make_blank_pdf(path: str, pages: int = 1) -> None:
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    with open(path, "wb") as f:
        writer.write(f)


def test_pdf_loader_raises_on_no_extractable_text(tmp_path):
    pdf_path = tmp_path / "blank.pdf"
    _make_blank_pdf(str(pdf_path), pages=2)
    with pytest.raises(NoExtractableTextError):
        PDFLoader().load(str(pdf_path))


def test_pdf_loader_raises_extraction_error_on_corrupt_file(tmp_path):
    bad_pdf = tmp_path / "bad.pdf"
    bad_pdf.write_bytes(b"this is not a real pdf")
    with pytest.raises(ExtractionError):
        PDFLoader().load(str(bad_pdf))


def test_docx_loader_extracts_paragraphs_and_headings(tmp_path):
    doc = docx.Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("This is the first paragraph of chapter one.")
    doc.add_paragraph("This is the second paragraph.")
    docx_path = tmp_path / "sample.docx"
    doc.save(str(docx_path))

    pages = DOCXLoader().load(str(docx_path))
    assert len(pages) == 1
    assert "Chapter One" in pages[0].text
    assert "first paragraph" in pages[0].text
    assert pages[0].section == "Chapter One"


def test_docx_loader_extracts_table_content(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Intro text")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Refund window"
    table.cell(1, 1).text = "30 days"
    docx_path = tmp_path / "table.docx"
    doc.save(str(docx_path))

    pages = DOCXLoader().load(str(docx_path))
    assert "Refund window" in pages[0].text
    assert "30 days" in pages[0].text


def test_docx_loader_raises_on_empty_document(tmp_path):
    doc = docx.Document()
    docx_path = tmp_path / "empty.docx"
    doc.save(str(docx_path))
    with pytest.raises(NoExtractableTextError):
        DOCXLoader().load(str(docx_path))


def test_txt_loader_reads_utf8_content(tmp_path):
    txt_path = tmp_path / "sample.txt"
    txt_path.write_text("Hello, world! Unicode: caf\u00e9, \u00e9l\u00e8ve.", encoding="utf-8")
    pages = TXTLoader().load(str(txt_path))
    assert "Hello, world!" in pages[0].text
    assert "caf\u00e9" in pages[0].text


def test_txt_loader_raises_on_empty_file(tmp_path):
    txt_path = tmp_path / "empty.txt"
    txt_path.write_text("", encoding="utf-8")
    with pytest.raises(NoExtractableTextError):
        TXTLoader().load(str(txt_path))


def test_txt_loader_normalizes_newlines(tmp_path):
    txt_path = tmp_path / "crlf.txt"
    txt_path.write_bytes(b"line one\r\nline two\r\n")
    pages = TXTLoader().load(str(txt_path))
    assert "\r" not in pages[0].text


def test_registry_returns_correct_loader_by_extension():
    assert isinstance(get_loader("pdf"), PDFLoader)
    assert isinstance(get_loader("docx"), DOCXLoader)
    assert isinstance(get_loader("txt"), TXTLoader)


def test_registry_raises_on_unsupported_extension():
    from app.core.exceptions import UnsupportedFileTypeError

    with pytest.raises(UnsupportedFileTypeError):
        get_loader("exe")
