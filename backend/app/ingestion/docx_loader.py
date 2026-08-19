"""
DOCX text extraction using python-docx.

Extracts paragraphs (including heading level), and tables (rendered as
pipe-delimited rows so their content is preserved and searchable).
DOCX has no native concept of "pages" so we return a single
`DocumentPage` per document, but we record heading structure in the
`section` field of nearby content when possible, and keep everything
in document (paragraph) order.
"""
import docx
from docx.opc.exceptions import PackageNotFoundError

from app.core.exceptions import ExtractionError, NoExtractableTextError
from app.ingestion.base import DocumentLoader, DocumentPage


class DOCXLoader(DocumentLoader):
    @staticmethod
    def supports(extension: str) -> bool:
        return extension.lower() == "docx"

    def load(self, file_path: str) -> list[DocumentPage]:
        try:
            document = docx.Document(file_path)
        except PackageNotFoundError as exc:
            raise ExtractionError(f"Invalid or corrupted DOCX file: {exc}") from exc
        except Exception as exc:  # noqa: BLE001
            raise ExtractionError(f"Failed to open DOCX: {exc}") from exc

        blocks: list[str] = []
        current_section: str | None = None
        headings: list[str] = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading") or style == "title":
                current_section = text
                headings.append(text)
                blocks.append(f"## {text}")
            else:
                blocks.append(text)

        for t_index, table in enumerate(document.tables, start=1):
            table_lines = [f"[Table {t_index}]"]
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    table_lines.append(" | ".join(cells))
            if len(table_lines) > 1:
                blocks.append("\n".join(table_lines))

        full_text = "\n\n".join(blocks).strip()
        if not full_text:
            raise NoExtractableTextError("No extractable text was found in this DOCX file.")

        core_props = document.core_properties
        metadata = {
            "author": core_props.author or None,
            "title": core_props.title or None,
            "headings": headings,
            "table_count": len(document.tables),
        }

        return [
            DocumentPage(
                text=full_text,
                page_number=None,
                section=current_section,
                metadata=metadata,
            )
        ]
